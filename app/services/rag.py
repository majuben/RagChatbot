"""
RAG service implementation connecting Ollama and PostgreSQL.

"""

import io
import os
from typing import List, Optional, Tuple
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaLLM as Ollama
from langchain_core.documents import Document
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import PGVector
import fitz
import docx

from .ollama import build_llm, build_embeddings

VECTORSTORE_TABLE_NAME = os.getenv("VECTORSTORE_TABLE_NAME")
DATABASE_URL = os.getenv("DATABASE_URL")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE"))       
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP"))  
TOP_K_RETRIEVAL = int(os.getenv("TOP_K_RETRIEVAL"))  
TOP_K_FINAL = int(os.getenv("TOP_K_FINAL"))           
SIMILARITY_THRESHOLD = float(os.getenv("SIMILARITY_THRESHOLD"))  
SUPPORTED_FILE_TYPES = [".pdf", ".docx"]


def _build_embeddings() -> OllamaEmbeddings:
    return build_embeddings()


def _build_llm() -> Ollama:
    return build_llm()


def _get_vectorstore() -> PGVector:
    embeddings = _build_embeddings()
    return PGVector(
        collection_name=VECTORSTORE_TABLE_NAME,
        connection_string=DATABASE_URL,
        embedding_function=embeddings
    )


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        return "\n\n".join(page.get_text() for page in document)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    return "\n\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def _extract_text_from_file(file_name: str, file_bytes: bytes) -> str:
    file_name_lower = file_name.lower()
    if file_name_lower.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    if file_name_lower.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")


def chunk_documents(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Découpe le texte en chunks avec des séparateurs sémantiques.
    Chunks plus petits = meilleure précision de retrieval.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", ", ", " ", ""],
        strip_whitespace=True,
        length_function=len,
    )
    chunks = splitter.split_text(text)
    return [c.strip() for c in chunks if c and len(c.strip()) > 20]



def deduplicate_chunks(
    docs_with_scores: List[Tuple[Document, float]],
    threshold: float = SIMILARITY_THRESHOLD
) -> List[Tuple[Document, float]]:
    """
    Supprime les chunks quasi-identiques basé sur le score de distance PGVector.
    """
    seen_previews = []
    deduplicated = []

    for doc, score in docs_with_scores:
        preview = doc.page_content[:120].strip()
        is_duplicate = False
        for seen in seen_previews:
            overlap_chars = sum(1 for a, b in zip(preview, seen) if a == b)
            similarity = overlap_chars / max(len(preview), len(seen), 1)
            if similarity > 0.75:
                is_duplicate = True
                break

        if not is_duplicate:
            seen_previews.append(preview)
            deduplicated.append((doc, score))

    return deduplicated


def retrieve_chunks(question: str, top_k: int = TOP_K_RETRIEVAL) -> List[Document]:
    """
    Récupère les chunks les plus pertinents avec déduplication.
    Retourne des Documents complets (avec metadata source) pour le regroupement.
    """
    vectorstore = _get_vectorstore()
    results_with_scores = vectorstore.similarity_search_with_score("search_query: " + question, k=top_k)

    print(f"\n--- DEBUG RETRIEVAL (Top {top_k}) ---")
    for i, (doc, score) in enumerate(results_with_scores):
        source = doc.metadata.get("source", "inconnu")
        print(f"Rank {i+1} | Score: {score:.4f} | Source: {source} | Preview: {doc.page_content[:70]}...")
    print("------------------------------------\n")
    deduplicated = deduplicate_chunks(results_with_scores)

    print(f"Après déduplication : {len(deduplicated)}/{len(results_with_scores)} chunks retenus")

    final_docs = [doc for doc, score in deduplicated[:TOP_K_FINAL]]
    return final_docs


def build_context(docs: List[Document]) -> str:
    """
    Regroupe les chunks par source pour donner une structure claire au LLM.
    """

    sources: dict[str, List[str]] = {}
    for doc in docs:
        source = doc.metadata.get("source", "Document sans nom")
        sources.setdefault(source, []).append(doc.page_content)


    context_parts = []
    for source_name, chunks in sources.items():
        combined = "\n\n".join(chunks)
        context_parts.append(f"[SOURCE : {source_name}]\n{combined}")

    return "\n\n===\n\n".join(context_parts)


def generate_response(question: str, context: str, sources: List[str]) -> str:
    """
    Génère une réponse complète en forçant la synthèse multi-sources.
    """
    llm = _build_llm()


    sources_list = "\n".join(f"- {s}" for s in set(sources)) if sources else "- Document unique"

    prompt = f"""Tu es un assistant expert en analyse documentaire. Tu disposes d'extraits issus de {len(set(sources))} source(s) :
{sources_list}

Ta mission : répondre à la question de façon COMPLÈTE et EXHAUSTIVE en exploitant TOUTES les informations pertinentes présentes dans le contexte, quelle que soit leur source.

RÈGLES STRICTES :
1. Synthétise et regroupe les informations complémentaires provenant de sources différentes.
2. Si plusieurs sources donnent des informations sur le même sujet, combine-les en une réponse unifiée.
3. Sois exhaustif : inclus tous les détails, conditions, étapes, valeurs numériques, exceptions.
4. Réponds directement et naturellement, sans mentionner "selon le document" ou "d'après le contexte".
5. Si une information n'est vraiment pas dans le contexte, dis-le brièvement à la fin.

CONTEXTE DOCUMENTAIRE :
{context}

QUESTION :
{question}

RÉPONSE COMPLÈTE ET SYNTHÉTISÉE :"""

    response = llm.invoke(prompt)
    return response.strip()


def ingest(texts: List[str], source: Optional[str] = None) -> None:
    all_documents = []
    for text in texts:
        chunks = chunk_documents(text)
        print(f"[INGEST] {len(chunks)} chunks générés pour '{source}'")
        for chunk in chunks:
            metadata = {"source": source} if source else {}
            all_documents.append(Document(
                page_content="search_document: " + chunk,
                metadata=metadata
            ))

    if all_documents:
        vectorstore = _get_vectorstore()
        
        existing = vectorstore.similarity_search(
            "test", k=1, filter={"source": source}
        )
        if existing:
            print(f"[INGEST] ⚠️ Source '{source}' déjà indexée, skip.")
            return
            
        vectorstore.add_documents(all_documents)
        print(f"[INGEST] ✅ {len(all_documents)} chunks indexés.")

    if all_documents:
        vectorstore = _get_vectorstore()
        vectorstore.add_documents(all_documents)
        print(f"[INGEST] ✅ {len(all_documents)} chunks indexés.")


def ingest_file(file_name: str, file_bytes: bytes) -> str:
    text = _extract_text_from_file(file_name, file_bytes)
    if not text or not text.strip():
        raise ValueError("No readable text found in uploaded file.")
    ingest([text], source=file_name)
    return text


def respond(question: str) -> str:
    """
    Pipeline RAG amélioré :
    1. Retrieval large (TOP_K_RETRIEVAL)
    2. Déduplication
    3. Sélection finale (TOP_K_FINAL)
    4. Regroupement par source
    5. Génération avec prompt de synthèse multi-sources
    """
    docs = retrieve_chunks(question, top_k=TOP_K_RETRIEVAL)

    if not docs:
        return "Je n'ai trouvé aucune information pertinente dans les documents indexés."

    context = build_context(docs)
    sources = [doc.metadata.get("source", "inconnu") for doc in docs]
    response = generate_response(question, context, sources)
    return response


class RAGService:
    @staticmethod
    def ingest(texts: List[str], source: Optional[str] = None) -> None:
        ingest(texts, source=source)

    @staticmethod
    def ingest_file(file_name: str, file_bytes: bytes) -> str:
        return ingest_file(file_name, file_bytes)

    @staticmethod
    def respond(question: str) -> str:
        return respond(question)


rag_service = RAGService()